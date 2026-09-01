"""
Assembles the final client letter as a formatted .docx.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build_letter_docx(
    client_name: str,
    portfolio_paragraph: str,
    macro_paragraph: str,
    recommendation_paragraph: str,
    advisor_name: str,
    output_path: str,
):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    doc.add_paragraph(f"Prezado {client_name},")
    intro = doc.add_paragraph(
        "Segue o relatório mensal referente ao desempenho do seu portfólio, ao cenário "
        "macroeconômico e as recomendações alinhadas ao seu perfil de investidor."
    )

    doc.add_heading("Desempenho do Portfólio", level=2)
    doc.add_paragraph(portfolio_paragraph)

    doc.add_heading("Cenario Macroeconômico", level=2)
    doc.add_paragraph(macro_paragraph)

    doc.add_heading("Recomendações", level=2)
    doc.add_paragraph(recommendation_paragraph)

    doc.add_paragraph(
        "Estamos à disposição para discutir os resultados e esclarecer qualquer dúvida."
    )
    doc.add_paragraph("Atenciosamente,")
    sign = doc.add_paragraph(advisor_name)
    sign.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save(output_path)
